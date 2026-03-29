from datetime import datetime
from docx import Document


def add_section(doc: Document, title: str, points: list[str]) -> None:
    doc.add_heading(title, level=1)
    for point in points:
        doc.add_paragraph(point, style="List Bullet")


def main() -> None:
    doc = Document()
    doc.add_heading("Online Voting System - Full Workflow", 0)
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    add_section(doc, "1. System Overview", [
        "This application is a Flask + SQLite based online voting platform with voter, admin, and owner roles.",
        "Voters can register, login, view active elections, cast one vote per election, and view results.",
        "Admins can create elections, add candidates, start/end elections, and monitor results.",
        "Owner (boss) has full control: can promote registered users to admin and transfer ownership.",
    ])

    add_section(doc, "2. Prerequisites", [
        "Python 3.10+ installed.",
        "Dependencies installed (Flask).",
        "Project files include app.py, schema.sql, templates folder, and voting.db.",
    ])

    add_section(doc, "3. Starting the Application", [
        "Open terminal in project root folder.",
        "Set owner setup key before first owner creation (PowerShell): $env:ADMIN_SETUP_KEY=\"YourStrongSecretKey\"",
        "Run: python app.py",
        "Open browser: http://127.0.0.1:5000",
    ])

    add_section(doc, "4. First-Time Owner (Boss) Setup Workflow", [
        "Register as a normal voter first using /register.",
        "Go to /admin/login and open First-Time Admin Setup.",
        "Enter registered voter email + voter password + owner setup key.",
        "If key and credentials match, account is created as role = owner.",
        "Login through admin login to access owner-enabled panel.",
    ])

    add_section(doc, "5. Admin Promotion Workflow (Owner Only)", [
        "Owner logs in and opens Admin Panel.",
        "Use Promote Registered User section.",
        "Enter email of an already registered voter.",
        "System copies that user into Admins table with role = admin.",
        "Promoted user can now login from admin login page.",
    ])

    add_section(doc, "6. Ownership Transfer Workflow (Owner Only)", [
        "Owner logs in and opens Transfer Ownership.",
        "Select target admin from the dropdown list.",
        "Confirm transfer action.",
        "Current owner becomes admin, selected admin becomes new owner.",
        "This creates a controlled new boss handover process.",
    ])

    add_section(doc, "7. Election Management Workflow", [
        "Admin/Owner creates election with title, start time, and end time.",
        "Validation ensures end time is after start time.",
        "Admin/Owner adds candidates linked to election.",
        "Admin/Owner starts election by setting status to active.",
        "Admin/Owner ends election by setting status to ended.",
    ])

    add_section(doc, "8. Voter Workflow", [
        "Voter registers and logs in.",
        "Dashboard shows active elections only.",
        "Voter opens election and selects one candidate.",
        "Vote is submitted once; duplicate voting is blocked by UNIQUE(user_id, election_id).",
        "Voter can view public results after voting period or as available.",
    ])

    add_section(doc, "9. Security and Integrity Controls", [
        "Password hashes are stored instead of plain text.",
        "Owner setup requires both registered credentials and secret setup key.",
        "Admin-only and owner-only routes are protected by decorators.",
        "Vote cast operation is wrapped in transaction with rollback on failure.",
        "Candidate-election relation is validated before vote insert.",
    ])

    add_section(doc, "10. Changing Owner Setup Key Anytime", [
        "Temporary for current terminal session: $env:ADMIN_SETUP_KEY=\"NewKey\"",
        "Permanent for Windows user: setx ADMIN_SETUP_KEY \"NewKey\"",
        "After setx, restart terminal and app to apply updated key.",
        "Keep this key private and share only with trusted authority.",
    ])

    add_section(doc, "11. Recommended Operational Policy", [
        "Keep one official owner account under faculty control.",
        "Promote student coordinators as admin only when required.",
        "Rotate owner setup key after each election cycle.",
        "Back up voting.db before major election operations.",
        "Do not share owner credentials in chat or screenshots.",
    ])

    add_section(doc, "12. Quick Troubleshooting", [
        "If admin setup fails: verify user is registered first and key is correct.",
        "If owner controls are missing: login with owner account, not admin account.",
        "If app does not start: install missing package and re-run python app.py.",
        "If key change not applying: restart terminal and server process.",
    ])

    output_path = "Voting_System_Full_Workflow.docx"
    doc.save(output_path)
    print(output_path)


if __name__ == "__main__":
    main()
